#!/usr/bin/env python3
"""
法律内容生产工具 - 客户告知书/进展说明函/法律问答

用法:
    python3 generate_content.py --type notice --input 通知内容.json --output 客户告知函.docx
    python3 generate_content.py --type guide --input 进展信息.json --output 进展说明函.docx
    python3 generate_content.py --type faq --input 问答内容.json --output 法律问答.docx

注: 普法文章（article）建议直接让 Claude 输出文本，无需脚本生成 Word。

内容类型:
    notice  - 客户告知书/通知函
    guide   - 案件进展说明函
    faq     - 法律问答整理
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


def set_font(run, name='仿宋_GB2312', size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_page_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, indent=False,
             font_name='仿宋_GB2312', size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_font(run, font_name, size, bold)
    return p


def generate_notice(doc, data):
    """生成客户告知书/通知函"""
    # 标题
    title = data.get('标题', '客户告知函')
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    run = title_p.add_run(title)
    set_font(run, '黑体', 18, bold=True)

    doc.add_paragraph()

    # 称谓
    recipient = data.get('收件方', '尊敬的当事人')
    add_para(doc, f'{recipient}：', size=12)

    # 正文
    content = data.get('正文内容', '')
    if content:
        for para in content.split('\n'):
            if para.strip():
                add_para(doc, para.strip(), indent=True)

    # 需要事项
    items = data.get('需要事项', [])
    if items:
        doc.add_paragraph()
        add_para(doc, '请您配合提供/办理以下事项：', indent=True)
        for i, item in enumerate(items, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.48)
            run = p.add_run(f'{i}. {item}')
            set_font(run, '仿宋_GB2312', 12)

    # 截止日期
    deadline = data.get('截止日期', '')
    if deadline:
        doc.add_paragraph()
        add_para(doc, f'请于 {deadline} 前完成上述事项，如有疑问请及时联系。', indent=True)

    doc.add_paragraph()

    # 联系方式
    contact = data.get('联系方式', '')
    if contact:
        add_para(doc, f'联系电话：{contact}')

    # 落款
    doc.add_paragraph()
    sender = data.get('发件方', '')
    law_firm = data.get('律所', '')
    date = data.get('日期', datetime.now().strftime('%Y年%m月%d日'))
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_text = ''
    if sender:
        sign_text += f'{sender}\n'
    if law_firm:
        sign_text += f'{law_firm}\n'
    sign_text += date
    run = sign_p.add_run(sign_text)
    set_font(run, '仿宋_GB2312', 12)


def generate_guide(doc, data):
    """生成案件进展说明函"""
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    run = title_p.add_run('案件进展说明函')
    set_font(run, '黑体', 18, bold=True)

    doc.add_paragraph()

    recipient = data.get('收件方', '尊敬的委托人')
    add_para(doc, f'{recipient}：')

    # 案件基本信息
    case_name = data.get('案件名称', '')
    if case_name:
        add_para(doc, f'您委托本所代理的"{case_name}"案件，现将进展情况报告如下：', indent=True)

    doc.add_paragraph()

    # 当前阶段
    stage = data.get('当前阶段', '')
    if stage:
        p = doc.add_paragraph()
        run = p.add_run('一、当前阶段')
        set_font(run, '黑体', 12, bold=True)
        add_para(doc, stage, indent=True)

    # 已完成工作
    completed = data.get('已完成工作', [])
    if completed:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run('二、已完成工作')
        set_font(run, '黑体', 12, bold=True)
        for i, item in enumerate(completed, 1):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(0.74)
            run = p2.add_run(f'{i}. {item}')
            set_font(run, '仿宋_GB2312', 12)

    # 下一步安排
    next_steps = data.get('下一步安排', [])
    if next_steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run('三、下一步安排')
        set_font(run, '黑体', 12, bold=True)
        for i, item in enumerate(next_steps, 1):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(0.74)
            run = p2.add_run(f'{i}. {item}')
            set_font(run, '仿宋_GB2312', 12)

    # 注意事项
    notes = data.get('注意事项', [])
    if notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run('四、注意事项')
        set_font(run, '黑体', 12, bold=True)
        for note in notes:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(0.74)
            run = p2.add_run(f'• {note}')
            set_font(run, '仿宋_GB2312', 12)

    doc.add_paragraph()
    add_para(doc, '如有任何疑问，请随时与我们联系。感谢您的信任与支持！', indent=True)

    # 落款
    doc.add_paragraph()
    sender = data.get('发件方', '')
    law_firm = data.get('律所', '')
    date = data.get('日期', datetime.now().strftime('%Y年%m月%d日'))
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_text = ''
    if sender:
        sign_text += f'经办律师：{sender}\n'
    if law_firm:
        sign_text += f'律师事务所：{law_firm}\n'
    sign_text += date
    run = sign_p.add_run(sign_text)
    set_font(run, '仿宋_GB2312', 12)


def generate_faq(doc, data):
    """生成法律问答整理"""
    title = data.get('标题', '法律问答整理')
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    run = title_p.add_run(title)
    set_font(run, '黑体', 18, bold=True)

    subtitle = data.get('副标题', '')
    if subtitle:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_p.add_run(subtitle)
        set_font(run, '仿宋_GB2312', 12)

    doc.add_paragraph()

    qa_list = data.get('问答列表', [])
    for i, qa in enumerate(qa_list, 1):
        # 问题
        q_p = doc.add_paragraph()
        q_p.paragraph_format.space_before = Pt(10)
        q_p.paragraph_format.space_after = Pt(4)
        q_run = q_p.add_run(f'Q{i}：{qa.get("问题", "")}')
        set_font(q_run, '黑体', 12, bold=True, color=RGBColor(0x17, 0x37, 0x6A))

        # 回答
        answer = qa.get('回答', '')
        if answer:
            for para in answer.split('\n'):
                if para.strip():
                    a_p = doc.add_paragraph()
                    a_p.paragraph_format.left_indent = Cm(0.74)
                    a_p.paragraph_format.space_after = Pt(4)
                    a_run = a_p.add_run(f'A：{para.strip()}' if para == answer else para.strip())
                    set_font(a_run, '仿宋_GB2312', 12)

        # 法律依据
        legal = qa.get('法律依据', '')
        if legal:
            legal_p = doc.add_paragraph()
            legal_p.paragraph_format.left_indent = Cm(0.74)
            run = legal_p.add_run(f'法律依据：{legal}')
            set_font(run, '仿宋_GB2312', 10, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()

    # 落款
    date = data.get('日期', datetime.now().strftime('%Y年%m月%d日'))
    law_firm = data.get('律所', '')
    if law_firm:
        sign_p = doc.add_paragraph()
        sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = sign_p.add_run(f'{law_firm}\n{date}')
        set_font(run, '仿宋_GB2312', 11)


def generate_content(data, content_type, output_path):
    """主生成函数"""
    doc = Document()
    set_page_margins(doc)

    doc.styles['Normal'].font.name = '仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    generators = {
        'notice': generate_notice,
        'guide': generate_guide,
        'faq': generate_faq,
    }

    if content_type not in generators:
        print(f"不支持的内容类型: {content_type}，支持: {list(generators.keys())}")
        sys.exit(1)

    generators[content_type](doc, data)

    doc.save(output_path)
    print(f"文档已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description='法律内容生产工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument('--type', required=True,
                        choices=['notice', 'guide', 'faq'],
                        help='内容类型: notice/guide/faq')
    parser.add_argument('--input', required=True, help='内容 JSON 文件路径')
    parser.add_argument('--output', default='法律内容.docx', help='输出 Word 文件路径')

    args = parser.parse_args()

    json_path = Path(args.input)
    if not json_path.exists():
        print(f"错误: JSON 文件不存在: {json_path}")
        sys.exit(1)

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    try:
        generate_content(data, args.type, args.output)
    except Exception as e:
        print(f"生成失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
