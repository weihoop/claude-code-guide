#!/usr/bin/env python3
"""
生成《法律AI技能使用手册》Word 文档
运行: python3 generate_help_doc.py
"""

from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    exit(1)


OUTPUT = Path(__file__).parent / "法律AI技能使用手册.docx"


def set_font(run, name='仿宋_GB2312', size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_page(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def h1(doc, text):
    """一级标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, '黑体', 16, bold=True, color=RGBColor(0x17, 0x37, 0x6A))
    # 下划线分隔
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run('─' * 38)
    run2.font.color.rgb = RGBColor(0x17, 0x37, 0x6A)
    run2.font.size = Pt(10)
    return p


def h2(doc, text):
    """二级标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, '黑体', 13, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
    return p


def h3(doc, text):
    """三级标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, '黑体', 12, bold=True)
    return p


def body(doc, text, indent=True):
    """正文段落"""
    for line in text.split('\n'):
        if not line.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        if indent:
            p.paragraph_format.first_line_indent = Pt(24)
        run = p.add_run(line.strip())
        set_font(run, '仿宋_GB2312', 12)
    return p


def tip_box(doc, text, title='提示', color='E8F4F8', border='2E74B5'):
    """提示框（带背景色的段落）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    t_run = p.add_run(f'【{title}】 ')
    set_font(t_run, '黑体', 11, bold=True, color=RGBColor(0x17, 0x37, 0x6A))
    c_run = p.add_run(text)
    set_font(c_run, '仿宋_GB2312', 11)
    return p


def bullet(doc, items, numbered=False):
    """项目列表"""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.space_after = Pt(4)
        prefix = f'{i}. ' if numbered else '• '
        run = p.add_run(prefix + item)
        set_font(run, '仿宋_GB2312', 12)


def code_block(doc, text):
    """代码块样式"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x24, 0x72, 0x2E)
    return p


def two_col_table(doc, headers, rows, col_widths=None, header_color='2F5496'):
    """双栏表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, '黑体', 11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for ri, row_data in enumerate(rows):
        fill = 'EBF1F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            set_cell_bg(cell, fill)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_font(run, '仿宋_GB2312', 11)

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)

    doc.add_paragraph()
    return table


def skill_card(doc, num, name, trigger, output, scenario):
    """技能卡片"""
    # 技能标题行
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    num_run = p.add_run(f'  {num}  ')
    set_font(num_run, '黑体', 13, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    # 修改段落背景——用表格模拟
    # 改用简单的粗体标题
    p.clear()
    run = p.add_run(f'▍ {num}. {name}')
    set_font(run, '黑体', 13, bold=True, color=RGBColor(0x17, 0x37, 0x6A))

    rows = [
        ['触发词', trigger],
        ['输出文件', output],
        ['典型场景', scenario],
    ]
    for label, val in rows:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(3)
        l_run = p2.add_run(f'{label}：')
        set_font(l_run, '黑体', 11, bold=True)
        v_run = p2.add_run(val)
        set_font(v_run, '仿宋_GB2312', 11)


def generate():
    doc = Document()
    set_page(doc)
    doc.styles['Normal'].font.name = '仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # ===================================================
    # 封面
    # ===================================================
    for _ in range(4):
        doc.add_paragraph()

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run('法律 AI 技能使用手册')
    set_font(run, '黑体', 28, bold=True, color=RGBColor(0x17, 0x37, 0x6A))

    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_sub.add_run('Claude Code × 律师日常工作场景')
    set_font(run, '仿宋_GB2312', 14, color=RGBColor(0x55, 0x55, 0x55))

    for _ in range(6):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f'版本：v1.0    更新日期：{datetime.now().strftime("%Y年%m月%d日")}')
    set_font(run, '仿宋_GB2312', 11, color=RGBColor(0x88, 0x88, 0x88))

    doc.add_page_break()

    # ===================================================
    # 第一章：基本概念
    # ===================================================
    h1(doc, '第一章  基本概念与工作原理')

    h2(doc, '1.1  三个核心名词解释')

    body(doc, '在开始使用之前，需要了解三个基础概念：终端（iTerm2）、Claude Code、以及 Skills 技能包。理解它们之间的关系，是高效使用这套工具的前提。')

    doc.add_paragraph()

    # 名词解释表格
    two_col_table(doc,
        ['名词', '通俗解释'],
        [
            ['终端\n（iTerm2）',
             '终端是一个"命令输入窗口"，类似于旧式电脑的黑色命令行界面。\n'
             'iTerm2 是 Mac 上最常用的终端软件，比系统自带的"终端"更好用。\n'
             '你在这里输入命令，电脑就按命令执行。\n'
             '快捷键打开：按 ⌘+空格 输入 iterm 回车'],
            ['Claude Code\n（claude 命令）',
             'Claude Code 是 Anthropic 公司出品的 AI 编程助手，安装后可以在终端里通过输入 claude 来启动。\n'
             '启动后，它就像一个非常聪明的助手坐在你旁边——你用自然语言说需求，它帮你完成工作。\n'
             '它能读文件、写文件、运行程序、分析文档，远不止聊天。'],
            ['Skills\n（技能包）',
             'Skills 是预先配置给 Claude Code 的"工作模板"，存放在 ~/.claude/skills/ 目录下。\n'
             '每个 skill 包含：触发词（Claude 识别后自动启用）、工作流程说明、Python 脚本（自动生成 Word/Excel）。\n'
             '类比：如果 Claude 是一位全能助手，Skills 就是你给助手定制的工作规范和专业工具箱。'],
        ],
        col_widths=[3.5, 12.5]
    )

    h2(doc, '1.2  工作原理——Claude + Skills 如何帮你完成工作')

    body(doc, '整个工作流程分为两个阶段：AI 分析阶段和脚本生成阶段。')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('阶段一：AI 分析阶段（Claude 负责）')
    set_font(run, '黑体', 12, bold=True)
    bullet(doc, [
        '你告诉 Claude 案件信息、合同内容或需要什么文书',
        'Claude 阅读 skill 的工作说明，按照法律专业逻辑进行分析',
        'Claude 输出结构化的 JSON 数据（一种机器可读的格式，类似填写好的表格）',
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('阶段二：脚本生成阶段（Python 脚本负责）')
    set_font(run, '黑体', 12, bold=True)
    bullet(doc, [
        'Python 脚本读取 Claude 输出的 JSON 数据',
        '按照预设的 Word/Excel 格式模板自动排版',
        '生成符合中国法院标准格式的正式文书文件',
    ])

    doc.add_paragraph()
    tip_box(doc,
        '你只需要负责"说清楚案件情况"，AI 负责分析，程序负责排版。整个过程中，你不需要懂任何编程知识。',
        '核心优势')

    h2(doc, '1.3  文件存放位置说明')

    two_col_table(doc,
        ['路径', '存放内容'],
        [
            ['~/.claude/skills/', '所有技能包的根目录（~ 代表用户主目录）'],
            ['~/.claude/skills/[技能名]/SKILL.md', '技能说明文件，Claude 读取后知道如何工作'],
            ['~/.claude/skills/[技能名]/scripts/', '生成 Word/Excel 的 Python 脚本'],
            ['~/.claude/skills/[技能名]/references/', '参考模板 JSON 文件，可直接复制使用'],
            ['~/.claude/skills/[技能名]/templates/', '诉讼文书 JSON 模板（legal-doc-draft 专用）'],
            ['~/Desktop/法律AI技能使用帮助文档/', '本手册及未来所有学习资料'],
        ],
        col_widths=[7, 9]
    )

    doc.add_page_break()

    # ===================================================
    # 第二章：快速开始
    # ===================================================
    h1(doc, '第二章  快速开始——从打开终端到生成文书')

    h2(doc, '2.1  第一步：打开终端并启动 Claude')

    body(doc, '每次使用前，按以下步骤打开工具：')
    bullet(doc, [
        '按 ⌘（Command）+ 空格键，打开 Spotlight 搜索',
        '输入 iterm，按回车键，打开 iTerm2 终端窗口',
        '在终端中输入以下命令并按回车：',
    ], numbered=True)

    code_block(doc, 'claude')

    bullet(doc, [
        '等待几秒，Claude 启动后会显示欢迎界面和提示符',
        '此时可以直接用中文输入你的需求',
    ], numbered=False)

    tip_box(doc,
        '如果输入 claude 后显示"command not found"，说明 Claude Code 尚未安装。请联系技术人员安装。',
        '常见问题')

    h2(doc, '2.2  第二步：告诉 Claude 你要做什么')

    body(doc, '在 Claude 提示符后，直接用中文描述你的需求。关键是要说出触发词，Claude 会自动识别并启用对应的技能包。')

    doc.add_paragraph()
    body(doc, '示例——起草劳动仲裁申请书：', indent=False)
    code_block(doc, '帮我起草一份劳动仲裁申请书，申请人张三，在XX公司工作了5年，被违法解除合同，工资15000元每月。')

    doc.add_paragraph()
    body(doc, '示例——计算工伤赔偿：', indent=False)
    code_block(doc, '帮我做工伤赔偿计算，8级伤残，月工资6000元，工作了4年。')

    doc.add_paragraph()
    body(doc, '示例——审查合同：', indent=False)
    code_block(doc, '帮我审查这份买卖合同（粘贴合同内容），我代表买方，重点关注付款条款和质量验收条款。')

    h2(doc, '2.3  第三步：保存 Claude 输出的 JSON 并运行脚本')

    body(doc, 'Claude 分析完成后，会输出一段 JSON 格式的文本（以 { 开始、以 } 结束的内容）。')

    bullet(doc, [
        '复制 Claude 输出的 JSON 内容',
        '新建一个文本文件，粘贴内容，保存为 .json 格式（如：案件信息.json）',
        '在终端中运行对应的 Python 脚本（Claude 会自动提示你运行哪个命令）',
        '在当前目录找到生成的 .docx 或 .xlsx 文件，用 Word 打开',
    ], numbered=True)

    tip_box(doc,
        'Claude 会自动帮你完成步骤2和3，你只需要确认命令并按回车。大部分情况下你不需要手动操作这些步骤。',
        '更简单的方式')

    doc.add_page_break()

    # ===================================================
    # 第三章：8个技能详细说明
    # ===================================================
    h1(doc, '第三章  8个法律技能详细使用说明')

    h2(doc, '3.1  技能总览')

    two_col_table(doc,
        ['编号', '技能名称', '触发词（说这些词就会启用）', '输出'],
        [
            ['1', '赔偿计算器', '赔偿计算 / 工伤赔偿 / 经济补偿金 / 劳动赔偿 / 交通赔偿', 'Excel'],
            ['2', '合同审查', '合同审查 / 审合同 / 合同风险 / 条款审查 / 合同把关', 'Word'],
            ['3', '诉讼文书起草', '起诉状 / 答辩状 / 上诉状 / 仲裁申请书 / 诉状', 'Word'],
            ['4', '证据整理', '证据目录 / 质证意见 / 整理证据 / 证据清单', 'Word'],
            ['5', '案情梳理与法律意见书', '案情梳理 / 法律意见 / 案件分析 / 法律意见书', 'Word'],
            ['6', '企业合规审查', '合规审查 / 合规风险 / 企业合规 / 劳动合规 / 用工合规', 'Word'],
            ['7', '法律内容生产', '普法文章 / 公众号 / 客户通知 / 客户告知书 / 法律科普', 'Word/文本'],
            ['8', '庭审代理意见书', '代理意见 / 庭审意见 / 出庭意见 / 代理词 / 辩护词', 'Word'],
        ],
        col_widths=[1, 3.5, 7.5, 2]
    )

    # ===================================================
    # 技能1
    # ===================================================
    h2(doc, '3.2  技能1：赔偿计算器（compensation-calc）')

    h3(doc, '支持的计算类型')
    two_col_table(doc,
        ['计算类型', '需要提供的信息', '计算结果'],
        [
            ['劳动争议经济补偿金', '月工资、工龄（年）、解除原因（违法/合法/协商）', 'N 或 2N，含月工资封顶计算'],
            ['工伤赔偿', '伤残等级（1-10级）、月工资、工龄', '一次性伤残补助金+医疗补助金+就业补助金'],
            ['交通事故赔偿', '城镇/农村、伤残等级、误工/护理/住院天数、医疗费', '各损失项目分项计算+精神损害抚慰金'],
            ['离婚财产分割', '共同财产清单（名称、估值、建议归属）', '平均分割方案+折价补偿金'],
        ],
        col_widths=[3, 6.5, 6.5]
    )

    h3(doc, '使用示例')
    body(doc, '方式一：直接告诉 Claude（Claude 会自动运行脚本）：', indent=False)
    code_block(doc, '帮我计算劳动赔偿：员工月工资12000元，工作了6年，公司违法解除合同。')
    code_block(doc, '帮我做工伤赔偿计算，7级伤残，月工资8000元，工作3年。')
    code_block(doc, '帮我算交通事故赔偿：城镇居民，10级伤残，误工30天，医疗费2万。')

    doc.add_paragraph()
    body(doc, '方式二：直接在终端运行脚本（适合参数精确场景）：', indent=False)
    code_block(doc, 'python3 ~/.claude/skills/compensation-calc/scripts/calc_compensation.py \\\n  --type labor --monthly-salary 12000 --years 6 --reason illegal --output 赔偿.xlsx')
    code_block(doc, 'python3 ~/.claude/skills/compensation-calc/scripts/calc_compensation.py \\\n  --type injury --monthly-salary 8000 --years 3 --disability-level 7 --output 工伤赔偿.xlsx')
    code_block(doc, 'python3 ~/.claude/skills/compensation-calc/scripts/calc_compensation.py \\\n  --type traffic --urban --disability-level 10 --lost-work-days 30 --medical-fee 20000 --output 交通赔偿.xlsx')

    h3(doc, '输出文件说明')
    bullet(doc, [
        'Sheet1（分项明细）：每一项赔偿的计算过程 + 法律依据',
        'Sheet2（汇总合计）：各项金额汇总 + 注意事项',
        'Sheet3（法条声明）：数据来源年份、AI生成提示、核实建议（flk.npc.gov.cn）',
    ])

    tip_box(doc,
        '标准参数（如全国平均工资）存放在 ~/.claude/skills/compensation-calc/references/标准参数.json，'
        '每年需手动更新一次（建议每年1月按文件内"_年度更新checklist"9步流程操作）。'
        '数据来源：国家统计局（data.stats.gov.cn）。',
        '重要提示')

    doc.add_paragraph()

    # ===================================================
    # 技能2
    # ===================================================
    h2(doc, '3.3  技能2：合同审查（contract-review）')

    h3(doc, '工作流程')
    bullet(doc, [
        '提供合同文本（直接粘贴全文）',
        '说明你的立场：甲方/乙方/中立',
        '说明合同类型（买卖/服务/劳动/房屋租赁等）',
        'Claude 逐条分析风险点，生成结构化结果',
        '运行脚本生成 Word 审查意见书',
    ], numbered=True)

    h3(doc, '输入示例')
    code_block(doc, '帮我审查这份合同（贴合同内容），我代表买方（甲方），这是一份买卖合同，重点关注付款条款和违约责任。')

    h3(doc, '输出格式')
    bullet(doc, [
        '封面（合同名称、审查立场、日期）',
        '总体评价',
        '风险清单表：条款位置 | 风险等级（高/中/低，带颜色）| 问题描述 | 修改建议',
        '缺失条款说明',
        '综合建议',
    ])

    h3(doc, '参考文件')
    bullet(doc, [
        'references/contract_review_template.json——审查结果 JSON 模板，可让 Claude 参考格式',
        'references/审查要点.md——8种合同类型的核查要点备忘录',
    ])

    doc.add_paragraph()

    # ===================================================
    # 技能3
    # ===================================================
    h2(doc, '3.4  技能3：诉讼文书起草（legal-doc-draft）')

    h3(doc, '支持的文书类型')
    two_col_table(doc,
        ['文书类型', '命令参数', '适用场景'],
        [
            ['民事起诉状', '--type qisu', '提起民事诉讼'],
            ['答辩状', '--type daben', '收到起诉状后 15 日内答辩'],
            ['上诉状', '--type susong', '一审判决后 15 日内上诉'],
            ['劳动仲裁申请书', '--type zhongcai', '申请劳动仲裁（仲裁前置）'],
            ['强制执行申请书', '--type zhixing', '胜诉判决后申请强制执行'],
        ],
        col_widths=[3.5, 3.5, 9]
    )

    h3(doc, '两种使用方式')
    body(doc, '方式一（推荐）：直接告诉 Claude，让其自动完成：', indent=False)
    code_block(doc, '帮我起草劳动仲裁申请书：申请人张三，被申请人XX公司，申请违法解除赔偿金150000元，工作5年月工资15000，2025年1月被违法解除合同。')

    doc.add_paragraph()
    body(doc, '方式二：填写模板 JSON 再生成：', indent=False)
    bullet(doc, [
        '从 templates/ 目录复制对应模板文件',
        '按模板格式填写案件信息',
        '运行脚本生成文书',
    ], numbered=True)

    h3(doc, '格式标准')
    bullet(doc, [
        '标题：黑体二号，居中',
        '正文：仿宋_GB2312 小四',
        '页面：A4，页边距 2.54cm',
        '完整落款：代理律师 + 律师事务所 + 日期',
    ])

    doc.add_paragraph()

    # ===================================================
    # 技能4
    # ===================================================
    h2(doc, '3.5  技能4：证据整理（evidence-organize）')

    h3(doc, '输出内容')
    bullet(doc, [
        '表1：己方证据目录——序号、证据名称、证明目的、来源、页码',
        '表2：对方证据质证意见——序号、证据名称、真实性/合法性/关联性（三性意见）、综合质证意见',
    ])

    h3(doc, '使用示例')
    code_block(doc, '帮我整理证据目录和质证意见：\n我方有3份证据（合同、发票、收货单）\n对方提交了2份证据（质检报告、退货申请单），我方对质检报告真实性有异议。')

    tip_box(doc,
        '可以直接告诉 Claude 每份证据的情况，Claude 会自动整理成标准的三性质证意见格式。',
        '技巧')

    doc.add_paragraph()

    # ===================================================
    # 技能5
    # ===================================================
    h2(doc, '3.6  技能5：案情梳理与法律意见书（case-summary）')

    h3(doc, '输出的法律意见书结构')
    bullet(doc, [
        '委托方基本信息',
        '案件事实概述',
        '法律关系分析',
        '争议焦点（分条列举）',
        '相关法律规定（引用具体条文）',
        '法律意见（针对每个争议焦点分条论述）',
        '证据建议清单',
        '风险提示',
        '律师声明 + 落款',
    ], numbered=True)

    h3(doc, '使用示例')
    code_block(doc, '帮我出具法律意见书：委托方李四，劳动争议案件，在某公司工作8年后被以"工作调整"为由降薪调岗，后公司以旷工为由解除合同，请分析法律关系并出具意见。')

    doc.add_paragraph()

    # ===================================================
    # 技能6
    # ===================================================
    h2(doc, '3.7  技能6：企业合规审查（compliance-check）')

    h3(doc, '审查领域')
    two_col_table(doc,
        ['参数', '审查领域', '核心内容'],
        [
            ['--area labor', '劳动用工合规', '劳动合同、规章制度、工时工资、社保公积金、解除程序'],
            ['--area data', '数据安全合规', '个人信息保护、数据分类分级、隐私政策（个人信息保护法）'],
            ['--area contract', '合同管理合规', '合同签订审批、履行管控、归档管理'],
            ['--area general', '综合合规', '以上所有领域综合审查'],
        ],
        col_widths=[2.5, 3.5, 10]
    )

    h3(doc, '输出格式')
    bullet(doc, [
        '审查概况（企业信息、审查范围）',
        '合规问题清单（表格，高/中/低风险颜色标注）',
        '整改建议（按优先级分组：高风险立即整改 → 中风险近期整改 → 低风险逐步优化）',
        '法律依据汇总',
    ])

    doc.add_paragraph()

    # ===================================================
    # 技能7
    # ===================================================
    h2(doc, '3.8  技能7：法律内容生产（legal-content）')

    h3(doc, '内容类型说明')
    two_col_table(doc,
        ['类型', '用途', '输出格式'],
        [
            ['普法文章（article）', '微信公众号法律科普文章，通俗易懂', '纯文本，可直接复制到公众号'],
            ['客户告知书（notice）', '通知当事人补材料、注意事项等', 'Word 文档'],
            ['进展说明函（guide）', '向当事人汇报案件进展情况', 'Word 文档'],
            ['法律问答（faq）', '整理当事人常见问题及解答', 'Word 文档'],
        ],
        col_widths=[4, 5.5, 6.5]
    )

    h3(doc, '使用示例')
    code_block(doc, '帮我写一篇公众号普法文章，主题是"试用期那些事儿"，1000字左右，通俗易懂。')
    code_block(doc, '帮我给当事人发一个告知函，告诉他需要在3月15日前补充提交银行流水和劳动合同原件。')
    code_block(doc, '帮我写一份案件进展说明，案件已开庭，等待判决，预计15个工作日出判决书。')

    doc.add_paragraph()

    # ===================================================
    # 技能8
    # ===================================================
    h2(doc, '3.9  技能8：庭审代理意见书（court-opinion）')

    h3(doc, '适用文书类型')
    bullet(doc, [
        '民事代理意见书（合同纠纷、侵权纠纷等）',
        '刑事辩护意见书（刑事案件）',
        '劳动仲裁代理意见书（劳动争议庭审）',
    ])

    h3(doc, '输出结构（符合中国法庭提交格式）')
    bullet(doc, [
        '标题：黑体二号居中',
        '代理人基本信息表（律师姓名、证号、律所、案号）',
        '第一部分：事实意见——己方事实陈述 + 对方主张反驳',
        '第二部分：证据意见——对庭审各方证据的综合意见',
        '第三部分：法律适用分析——引用法条 + 法律分析',
        '请求事项（列明具体诉求）',
        '落款（律师签名 + 律所 + 日期）',
    ], numbered=True)

    h3(doc, '使用示例')
    code_block(doc, '帮我写庭审代理意见书，我代理原告（买方），对方（卖方）拒不付货款80万，庭审中对方提出货物有质量问题，我方认为其质检报告不合法，请帮我整理代理意见。')

    tip_box(doc,
        '代理意见书是出庭最核心的文书，建议在庭审前一天完成，庭审当天可根据实际情况口头补充。',
        '实践建议')

    doc.add_page_break()

    # ===================================================
    # 第四章：完整工作流示例
    # ===================================================
    h1(doc, '第四章  完整工作流示例')

    h2(doc, '4.1  示例：劳动争议案件从接案到庭审的全流程')

    body(doc, '以下是一个完整的劳动争议案件处理流程，展示如何综合使用多个技能：')

    steps = [
        ('接案阶段——法律意见书', 'case-summary（案情梳理）',
         '告诉 Claude：委托方张三，在某公司工作5年被违法解除，工资15000元，\n帮我出具法律意见书，分析胜诉可能性和赔偿金额。'),
        ('申请仲裁——劳动仲裁申请书', 'legal-doc-draft（文书起草）',
         '帮我起草劳动仲裁申请书，申请人张三，被申请人XX公司，\n申请违法解除赔偿金150000元和拖欠工资60000元。'),
        ('证据准备——证据目录和质证意见', 'evidence-organize（证据整理）',
         '帮我整理证据目录，我方有劳动合同、工资流水、解除通知书3份证据。\n对方提交了规章制度和考勤记录，帮我准备质证意见。'),
        ('赔偿核算——确认金额', 'compensation-calc（赔偿计算）',
         '帮我计算劳动赔偿：月工资15000元，工龄5年，违法解除，输出Excel。'),
        ('庭审准备——代理意见书', 'court-opinion（庭审代理意见）',
         '帮我写庭审代理意见书，我代理申请人张三，庭审中对方主张\n申请人旷工属实，我方认为公司未证明录用条件，请帮我整理代理意见。'),
        ('庭后——告知当事人进展', 'legal-content（法律内容）',
         '帮我写一份案件进展说明函给当事人，告知今日已庭审完毕，\n等待仲裁裁决，预计45个工作日。'),
    ]

    for i, (phase, skill_name, example) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f'步骤 {i}：{phase}')
        set_font(run, '黑体', 12, bold=True, color=RGBColor(0x17, 0x37, 0x6A))

        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(2)
        run2 = p2.add_run(f'使用技能：{skill_name}')
        set_font(run2, '仿宋_GB2312', 11, color=RGBColor(0x55, 0x55, 0x55))

        code_block(doc, example)

    doc.add_page_break()

    # ===================================================
    # 第五章：常见问题
    # ===================================================
    h1(doc, '第五章  常见问题解答')

    qa_list = [
        ('Q：如何判断 Claude 是否识别了正确的技能？',
         'A：Claude 启动对应技能后，通常会在回复开头说明"我将使用 XXX 技能"或描述工作步骤。如果没有按预期工作，可以明确说：请使用劳动仲裁申请书技能帮我起草文书。'),
        ('Q：生成的 Word 文件乱码怎么办？',
         'A：确保用 Microsoft Word（Mac 版）打开，不要用其他软件。如果还是乱码，可以尝试：右键文件→打开方式→Microsoft Word。'),
        ('Q：计算结果和实际裁判结果不一致怎么办？',
         'A：计算结果基于全国平均标准参数，实际裁判以当地（省/市）标准为准，且法官有一定自由裁量空间。建议将计算结果作为参考，在标准参数.json 中更新为你所在省份的数据。'),
        ('Q：可以同时处理多件案子吗？',
         'A：可以。建议为每件案子建立独立的工作文件夹，在对应文件夹中生成文书，避免文件混淆。'),
        ('Q：Claude 的回答是否具有法律效力？',
         'A：没有。Claude 是辅助工具，所有输出的文书需要由执业律师审核确认后才能使用。Claude 可以帮你快速起草初稿，但专业判断和最终把关仍需律师本人负责。'),
        ('Q：如何更新标准参数（如全国平均工资）？',
         'A：打开 ~/.claude/skills/compensation-calc/references/标准参数.json，按照文件中"_年度更新checklist"的9步流程操作。'
         '主要数据来源：国家统计局官网（data.stats.gov.cn），查找"城镇非私营单位年均工资"和"居民人均可支配收入"两项指标。'
         '更新完 JSON 后，还需同步修改 calc_compensation.py 中 DEFAULT_PARAMS 字典的对应数值。'),
        ('Q：法律意见书/代理意见书里的法条是否可靠？',
         'A：系统已内置三层保障机制：\n'
         '①本地法条白名单（~/.claude/skills/shared/references/法条库.json）：收录25条高频法条原文，'
         'Claude 优先从此库精确匹配，标注"[库中已核实]"；\n'
         '②三级确定性标注：库中没有的条文，Claude 会标注"[待核实]"，不会编造条文号；\n'
         '③废止法律黑名单：系统禁止引用婚姻法、合同法、侵权责任法等已废止法律。\n'
         '即便如此，AI 引用法条仍可能存在偏差，重要文书使用前建议在国家法律法规数据库（flk.npc.gov.cn）核实。'),
        ('Q：生成的文书末尾有"法条可靠性声明"灰色小字，可以删掉吗？',
         'A：可以。那段灰色小字是自动附加的风险提示，不影响文书主体内容。'
         '如需提交正式文书，使用 Word 手动删除该段落即可。'),
        ('Q：法条库中没有我需要的法条怎么办？',
         'A：系统默认按三级协议处理——Claude 会在条文末标注"[待核实]"并注明条文号，'
         '供你手动查证后确认。如需将常用法条加入白名单，'
         '可编辑 ~/.claude/skills/shared/references/法条库.json 按现有格式添加。'),
        ('Q：案例参考库是什么？如何使用？',
         'A：系统已内置中国法院2024-2025年度案例库（约1739真实判例，覆盖23个法律领域）。\n'
         '使用案情梳理、代理意见书、合同审查、合规审查、文书起草等技能时，Claude 会自动检索相似案例供参考。\n'
         '案例索引文件位于 ~/.claude/skills/shared/references/案例索引.json，'
         '完整案例文本存储在 案例全文/ 目录下。\n'
         '注意：案例仅供参考，不构成判决依据。各地法院裁判尺度可能不同。'),
        ('Q：如何更新案例参考库？',
         'A：当获取新年度的法院案例PDF后，运行提取脚本即可更新：\n'
         'python3 ~/.claude/skills/shared/scripts/build_case_index.py '
         '--input-dir <PDF目录> --year <年份> --output <输出索引.json>\n'
         '然后用 --merge 参数合并新旧索引。详见脚本的 --help 说明。'),
    ]

    for q, a in qa_list:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        q_run = p.add_run(q)
        set_font(q_run, '黑体', 12, bold=True, color=RGBColor(0x17, 0x37, 0x6A))

        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(4)
        a_run = p2.add_run(a)
        set_font(a_run, '仿宋_GB2312', 12)

    doc.add_page_break()

    # ===================================================
    # 第六章：后续优化方向
    # ===================================================
    h1(doc, '第六章  后续优化方向')

    h2(doc, '6.1  短期优化（可自行完成）')

    two_col_table(doc,
        ['优化项目', '操作方法'],
        [
            ['更新赔偿计算标准参数',
             '每年更新 compensation-calc/references/标准参数.json 中的\n全国/省级平均工资数据'],
            ['添加常用律所信息',
             '在各 references/ 模板 JSON 中预填律师姓名、律所名称、联系方式，\n避免每次重复输入'],
            ['积累常用案件模板',
             '将常见案件类型（劳动争议、买卖合同、借贷纠纷）的完整 JSON 模板保存，\n下次直接复制修改'],
            ['添加更多合同审查要点',
             '在 contract-review/references/审查要点.md 中补充你在实践中\n发现的行业特殊审查要点'],
        ],
        col_widths=[4.5, 11.5]
    )

    h2(doc, '6.2  中期扩展（需少量技术支持）')

    bullet(doc, [
        '案件管理整合：将 8 个技能与案件档案系统整合，一个案号自动关联所有文书',
        '批量处理：对于批量类似案件（如群体劳动争议），支持一次性生成多份文书',
        '省级参数库：在标准参数.json 中补充全部 31 个省市的社平工资，自动适配当地标准',
        '文书版本管理：保留每次生成的文书版本，方便对比修改历史',
        '微信/邮件集成：客户告知书生成后直接发送到客户微信或邮箱',
    ])

    h2(doc, '6.3  已完成的优化项目')

    two_col_table(doc,
        ['优化项目', '完成状态', '说明'],
        [
            ['法条白名单库', '已完成',
             '25条高频法条原文入库，Claude 优先精确匹配'],
            ['案例参考库', '已完成',
             '中国法院2024-2025年度案例（1739判例，23个领域）已提取索引，\n'
             '5个核心 skill 已集成案例参考指令'],
        ],
        col_widths=[3.5, 2, 10.5]
    )

    h2(doc, '6.4  长期方向（平台化建设）')

    bullet(doc, [
        '案件数据库：建立本所案件数据库，历史相似案件可作为参考',
        '裁判结果预测：基于历史数据对案件胜诉率、赔偿金额进行预测',
        '法规自动更新：当相关法律法规修订时，自动更新计算参数和审查要点',
        '多律师协作：支持团队多律师共同处理同一案件的文书流转',
        '客户端小程序：当事人通过小程序自行填写信息，律师直接调用生成文书',
    ])

    h2(doc, '6.5  本帮助文档的更新计划')

    body(doc, '本手册存放于桌面"法律AI技能使用帮助文档"文件夹，后续所有更新均在此文件夹中进行。')

    bullet(doc, [
        '每新增一个技能，补充对应章节的使用说明',
        '遇到常见问题，在第五章补充解答',
        '每年1月，更新赔偿计算标准参数并在本文档中备注',
    ])

    # ===================================================
    # 附录
    # ===================================================
    doc.add_page_break()
    h1(doc, '附录  快速参考卡片')

    h2(doc, '附录A：技能触发词速查')

    two_col_table(doc,
        ['说这些词……', '启用的技能', '生成文件类型'],
        [
            ['赔偿计算、工伤赔偿、经济补偿金、劳动赔偿、交通赔偿', '赔偿计算器', 'Excel'],
            ['合同审查、审合同、合同风险、条款审查、合同把关', '合同审查', 'Word'],
            ['起诉状、答辩状、上诉状、仲裁申请书、诉状', '诉讼文书起草', 'Word'],
            ['证据目录、质证意见、整理证据、证据清单', '证据整理', 'Word'],
            ['案情梳理、法律意见、案件分析、法律意见书', '法律意见书', 'Word'],
            ['合规审查、合规风险、企业合规、劳动合规', '企业合规审查', 'Word'],
            ['普法文章、公众号、客户通知、客户告知书', '法律内容生产', 'Word/文本'],
            ['代理意见、庭审意见、代理词、辩护词', '庭审代理意见书', 'Word'],
        ],
        col_widths=[7, 3.5, 2.5]
    )

    h2(doc, '附录B：文件存放路径')

    two_col_table(doc,
        ['文件', '路径'],
        [
            ['技能包根目录', '~/.claude/skills/'],
            ['赔偿标准参数', '~/.claude/skills/compensation-calc/references/标准参数.json'],
            ['合同审查模板', '~/.claude/skills/contract-review/references/contract_review_template.json'],
            ['合同审查要点备忘', '~/.claude/skills/contract-review/references/审查要点.md'],
            ['诉讼文书模板（5种）', '~/.claude/skills/legal-doc-draft/templates/'],
            ['证据整理模板', '~/.claude/skills/evidence-organize/references/evidence_template.json'],
            ['法律意见书模板', '~/.claude/skills/case-summary/references/opinion_template.json'],
            ['合规核查清单', '~/.claude/skills/compliance-check/references/compliance_checklist.json'],
            ['合规审查模板', '~/.claude/skills/compliance-check/references/compliance_template.json'],
            ['客户告知书模板', '~/.claude/skills/legal-content/references/notice_template.json'],
            ['进展说明函模板', '~/.claude/skills/legal-content/references/guide_template.json'],
            ['代理意见书模板', '~/.claude/skills/court-opinion/references/court_opinion_template.json'],
            ['法条白名单库（新）', '~/.claude/skills/shared/references/法条库.json'],
            ['案例索引（总索引）', '~/.claude/skills/shared/references/案例索引.json'],
            ['案例索引（2024年）', '~/.claude/skills/shared/references/案例索引-2024.json'],
            ['案例索引（2025年）', '~/.claude/skills/shared/references/案例索引-2025.json'],
            ['案例全文目录', '~/.claude/skills/shared/references/案例全文/{年份}/{册号}-{领域}/'],
            ['案例提取脚本', '~/.claude/skills/shared/scripts/build_case_index.py'],
            ['本手册', '~/Desktop/Law-AI-帮助文档/法律AI技能使用手册.docx'],
        ],
        col_widths=[4.5, 11.5]
    )

    h2(doc, '附录C：法条可靠性保障机制快速参考')

    body(doc, '系统已内置四层法条可靠性防护，以下是各层机制的说明：')

    two_col_table(doc,
        ['防护层', '机制说明', '效果'],
        [
            ['第一层\nPrompt 约束',
             '5个核心 skill（案情梳理/代理意见/合同审查/合规审查/文书起草）均已配置"法条引用规则"。\n'
             'Claude 被要求：禁止引用废止法律、禁止编造条文号、对不确定条文标注[待核实]',
             '防止引用过期法条、防止编造'],
            ['第二层\n法条白名单库',
             '本地存储25条高频法条原文（劳动合同法、工伤保险条例、民法典等）。\n'
             'Claude 优先精确匹配原文，标注[库中已核实]。\n'
             '文件位置：~/.claude/skills/shared/references/法条库.json',
             '最高可靠性，原文引用'],
            ['第三层\n数据错误修复',
             '已修正赔偿计算器中城镇居民人均可支配收入数值（49283元/年，2023年）\n'
             '与标准参数.json 保持一致，消除两套数据并存的矛盾。',
             '计算数值准确'],
            ['第四层\n参数数据治理',
             '标准参数.json 新增：_权威来源（国家统计局查询路径）、\n'
             '_年度更新checklist（9步更新流程）、_内部一致性检查（JSON与Python同步清单）。',
             '年度维护有据可查'],
            ['第五层\n文档声明',
             '所有生成的 Word 文书末尾自动附加灰色小字法条声明；\n'
             'Excel 赔偿计算表新增 Sheet3"法条声明"，注明数据来源年份和核实建议。',
             '提示用户核实'],
        ],
        col_widths=[2.5, 10, 3.5]
    )

    doc.add_paragraph()
    tip_box(doc,
        '黑名单（禁止引用）：婚姻法→民法典婚姻家庭编、合同法→民法典合同编、'
        '侵权责任法→民法典侵权责任编、物权法→民法典物权编、继承法→民法典继承编。\n'
        '核实建议：国家法律法规数据库 flk.npc.gov.cn（全文免费查询）',
        '已废止法律黑名单')

    # 落款
    doc.add_paragraph()
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_p.add_run(f'—— 文档生成于 {datetime.now().strftime("%Y年%m月%d日")} ——')
    set_font(run, '仿宋_GB2312', 11, color=RGBColor(0x88, 0x88, 0x88))

    doc.save(OUTPUT)
    print(f'手册已生成: {OUTPUT}')


if __name__ == '__main__':
    generate()
