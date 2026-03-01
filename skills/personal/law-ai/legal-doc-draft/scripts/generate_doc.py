#!/usr/bin/env python3
"""
诉讼文书生成器 - 起诉状/答辩状/上诉状/仲裁申请书/强制执行申请书

用法:
    python3 generate_doc.py --type qisu --input 案件信息.json --output 民事起诉状.docx
    python3 generate_doc.py --type daben --input 案件信息.json --output 答辩状.docx
    python3 generate_doc.py --type susong --input 案件信息.json --output 上诉状.docx
    python3 generate_doc.py --type zhongcai --input 案件信息.json --output 仲裁申请书.docx
    python3 generate_doc.py --type zhixing --input 案件信息.json --output 执行申请书.docx

支持文书类型:
    qisu     - 民事起诉状
    daben    - 答辩状
    susong   - 上诉状
    zhongcai - 劳动仲裁申请书
    zhixing  - 强制执行申请书
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


# 文书标题映射
DOC_TITLES = {
    'qisu': '民事起诉状',
    'daben': '答辩状',
    'susong': '上诉状',
    'zhongcai': '劳动仲裁申请书',
    'zhixing': '强制执行申请书',
}

# 收件方映射
DOC_RECIPIENT = {
    'qisu': '人民法院',
    'daben': '人民法院',
    'susong': '人民法院',
    'zhongcai': '劳动争议仲裁委员会',
    'zhixing': '人民法院执行局',
}


def set_font(run, name='仿宋_GB2312', size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_page_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_title(doc, text, size=22):
    """添加文书标题（黑体二号居中）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_font(run, '黑体', size, bold=True)
    return p


def add_section_title(doc, text):
    """添加章节标题（黑体小三）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, '黑体', 14, bold=True)
    return p


def add_body_para(doc, text, indent=0):
    """添加正文段落（仿宋小四）"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24 * indent) if indent else Pt(24)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, '仿宋_GB2312', 12)
    return p


def add_party_table(doc, parties):
    """添加当事人信息表"""
    if not parties:
        return

    # 计算所有当事人的字段最大数
    all_fields = set()
    for party in parties:
        all_fields.update(party.keys())

    standard_fields = ['身份', '姓名/名称', '性别', '出生日期', '民族',
                        '住所地', '联系电话', '法定代表人', '统一社会信用代码']
    fields = [f for f in standard_fields if f in all_fields]
    fields += [f for f in all_fields if f not in standard_fields]

    for party in parties:
        role = party.get('身份', '当事人')
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        role_run = p.add_run(f'{role}：')
        set_font(role_run, '黑体', 12, bold=True)

        for field in fields:
            if field == '身份':
                continue
            val = party.get(field)
            if val:
                item_p = doc.add_paragraph()
                item_p.paragraph_format.left_indent = Cm(0.74)
                item_p.paragraph_format.space_after = Pt(3)
                run = item_p.add_run(f'{field}：{val}')
                set_font(run, '仿宋_GB2312', 12)


def generate_qisu(doc, data):
    """生成民事起诉状"""
    court = data.get('受理法院', '______人民法院')

    # 抬头
    p = doc.add_paragraph()
    run = p.add_run(f'致：{court}')
    set_font(run, '仿宋_GB2312', 12)

    doc.add_paragraph()

    # 当事人信息
    add_section_title(doc, '当事人信息')
    add_party_table(doc, data.get('当事人', []))
    doc.add_paragraph()

    # 诉讼请求
    add_section_title(doc, '诉讼请求')
    requests = data.get('诉讼请求', [])
    for i, req in enumerate(requests, 1):
        add_body_para(doc, f'{i}. {req}', indent=0)
    doc.add_paragraph()

    # 事实与理由
    add_section_title(doc, '事实与理由')
    facts = data.get('事实与理由', '')
    for para in facts.split('\n'):
        if para.strip():
            add_body_para(doc, para.strip())
    doc.add_paragraph()

    # 证据清单
    evidence = data.get('证据清单', [])
    if evidence:
        add_section_title(doc, '证据清单')
        for i, ev in enumerate(evidence, 1):
            add_body_para(doc, f'{i}. {ev}', indent=0)
        doc.add_paragraph()


def generate_daben(doc, data):
    """生成答辩状"""
    court = data.get('受理法院', '______人民法院')

    p = doc.add_paragraph()
    run = p.add_run(f'致：{court}')
    set_font(run, '仿宋_GB2312', 12)

    doc.add_paragraph()

    # 当事人信息
    add_section_title(doc, '当事人信息')
    add_party_table(doc, data.get('当事人', []))
    doc.add_paragraph()

    # 答辩意见
    add_section_title(doc, '答辩意见')
    opinion_p = doc.add_paragraph()
    run = opinion_p.add_run('答辩人就原告的诉讼请求，提出如下答辩意见：')
    set_font(run, '仿宋_GB2312', 12)

    opinions = data.get('答辩意见', [])
    for i, op in enumerate(opinions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        run = p.add_run(f'{i}. {op}')
        set_font(run, '仿宋_GB2312', 12)
    doc.add_paragraph()

    # 事实与理由
    add_section_title(doc, '事实与理由')
    facts = data.get('事实与理由', '')
    for para in facts.split('\n'):
        if para.strip():
            add_body_para(doc, para.strip())
    doc.add_paragraph()


def generate_susong(doc, data):
    """生成上诉状"""
    court = data.get('上诉法院', '______中级人民法院')

    p = doc.add_paragraph()
    run = p.add_run(f'致：{court}')
    set_font(run, '仿宋_GB2312', 12)

    doc.add_paragraph()

    add_section_title(doc, '当事人信息')
    add_party_table(doc, data.get('当事人', []))
    doc.add_paragraph()

    # 原审信息
    original_info = data.get('原审信息', {})
    if original_info:
        add_section_title(doc, '原审信息')
        for key, val in original_info.items():
            p = doc.add_paragraph()
            run = p.add_run(f'{key}：{val}')
            set_font(run, '仿宋_GB2312', 12)
        doc.add_paragraph()

    # 上诉请求
    add_section_title(doc, '上诉请求')
    for i, req in enumerate(data.get('上诉请求', []), 1):
        add_body_para(doc, f'{i}. {req}', indent=0)
    doc.add_paragraph()

    # 上诉理由
    add_section_title(doc, '上诉理由')
    reasons = data.get('上诉理由', '')
    for para in reasons.split('\n'):
        if para.strip():
            add_body_para(doc, para.strip())
    doc.add_paragraph()


def generate_zhongcai(doc, data):
    """生成劳动仲裁申请书"""
    committee = data.get('仲裁委', '______劳动争议仲裁委员会')

    p = doc.add_paragraph()
    run = p.add_run(f'致：{committee}')
    set_font(run, '仿宋_GB2312', 12)

    doc.add_paragraph()

    add_section_title(doc, '当事人信息')
    add_party_table(doc, data.get('当事人', []))
    doc.add_paragraph()

    add_section_title(doc, '仲裁请求')
    for i, req in enumerate(data.get('仲裁请求', []), 1):
        add_body_para(doc, f'{i}. {req}', indent=0)
    doc.add_paragraph()

    add_section_title(doc, '事实与理由')
    facts = data.get('事实与理由', '')
    for para in facts.split('\n'):
        if para.strip():
            add_body_para(doc, para.strip())
    doc.add_paragraph()

    evidence = data.get('证据清单', [])
    if evidence:
        add_section_title(doc, '证据清单')
        for i, ev in enumerate(evidence, 1):
            add_body_para(doc, f'{i}. {ev}', indent=0)
        doc.add_paragraph()


def generate_zhixing(doc, data):
    """生成强制执行申请书"""
    court = data.get('执行法院', '______人民法院')

    p = doc.add_paragraph()
    run = p.add_run(f'致：{court}执行局')
    set_font(run, '仿宋_GB2312', 12)

    doc.add_paragraph()

    add_section_title(doc, '当事人信息')
    add_party_table(doc, data.get('当事人', []))
    doc.add_paragraph()

    # 生效文书信息
    judgment = data.get('生效文书', {})
    if judgment:
        add_section_title(doc, '申请执行依据')
        for key, val in judgment.items():
            p = doc.add_paragraph()
            run = p.add_run(f'{key}：{val}')
            set_font(run, '仿宋_GB2312', 12)
        doc.add_paragraph()

    add_section_title(doc, '申请执行事项')
    for i, item in enumerate(data.get('执行事项', []), 1):
        add_body_para(doc, f'{i}. {item}', indent=0)
    doc.add_paragraph()

    # 被执行人财产线索
    property_clues = data.get('财产线索', [])
    if property_clues:
        add_section_title(doc, '被执行人财产线索')
        for clue in property_clues:
            add_body_para(doc, f'• {clue}', indent=0)
        doc.add_paragraph()


def generate_doc(data, doc_type, output_path):
    """主生成函数"""
    doc = Document()
    set_page_margins(doc)

    doc.styles['Normal'].font.name = '仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 文书标题
    title = data.get('文书标题', DOC_TITLES.get(doc_type, '法律文书'))
    add_title(doc, title)

    # 根据类型生成内容
    generators = {
        'qisu': generate_qisu,
        'daben': generate_daben,
        'susong': generate_susong,
        'zhongcai': generate_zhongcai,
        'zhixing': generate_zhixing,
    }
    generators[doc_type](doc, data)

    # 此致
    recipient = data.get('此致', DOC_RECIPIENT.get(doc_type, ''))
    if recipient:
        ci_p = doc.add_paragraph()
        ci_run = ci_p.add_run('此致')
        set_font(ci_run, '仿宋_GB2312', 12)

        court_p = doc.add_paragraph()
        court_p.paragraph_format.first_line_indent = Cm(0.74)
        court_run = court_p.add_run(recipient)
        set_font(court_run, '仿宋_GB2312', 12)

        doc.add_paragraph()

    # 落款
    date = data.get('日期', datetime.now().strftime('%Y年%m月%d日'))
    applicant = data.get('申请人', data.get('答辩人', data.get('原告', data.get('上诉人', ''))))
    agent = data.get('代理律师', '')
    law_firm = data.get('律师事务所', '')

    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_text = ''
    if applicant:
        sign_text += f'申请人/当事人：{applicant}\n'
    if agent:
        sign_text += f'代理律师：{agent}\n'
    if law_firm:
        sign_text += f'律师事务所：{law_firm}\n'
    sign_text += f'日期：{date}'
    sign_run = sign_p.add_run(sign_text)
    set_font(sign_run, '仿宋_GB2312', 12)

    doc.save(output_path)
    print(f"{title}已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description='诉讼文书生成器',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument('--type', required=True,
                        choices=['qisu', 'daben', 'susong', 'zhongcai', 'zhixing'],
                        help='文书类型')
    parser.add_argument('--input', required=True, help='案件信息 JSON 文件路径')
    parser.add_argument('--output', default='诉讼文书.docx', help='输出 Word 文件路径')

    args = parser.parse_args()

    json_path = Path(args.input)
    if not json_path.exists():
        print(f"错误: JSON 文件不存在: {json_path}")
        sys.exit(1)

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    try:
        generate_doc(data, args.type, args.output)
    except Exception as e:
        print(f"生成失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
