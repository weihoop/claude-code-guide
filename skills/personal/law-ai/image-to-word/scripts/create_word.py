#!/usr/bin/env python3
"""
JSON转Word工具 - 带格式美化

用法:
    python3 create_word.py <json文件路径> [输出docx路径]

示例:
    python3 create_word.py data.json
    python3 create_word.py data.json output.docx

JSON格式说明:
{
  "title": "文档标题",
  "sections": [
    {
      "heading": "章节标题",
      "type": "table",           # table/text/list
      "headers": ["列1", "列2"],  # 表格表头
      "rows": [["数据1", "数据2"]],  # 表格数据
      "header_color": "FFF2CC"   # 可选，表头颜色
    },
    {
      "type": "text",
      "content": "段落内容"
    },
    {
      "type": "list",
      "items": ["项目1", "项目2"]
    }
  ]
}
"""

import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


def set_cell_shading(cell, color: str):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_table(doc, headers: list, rows: list, header_color: str = 'D9E2F3'):
    """
    创建格式化表格

    Args:
        doc: Document对象
        headers: 表头列表
        rows: 数据行列表
        header_color: 表头背景色（十六进制，不带#）
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        para = header_cells[i].paragraphs[0]
        if para.runs:
            para.runs[0].bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], header_color)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = str(cell_data)
                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def json_to_word(json_path: str, docx_path: str = None) -> str:
    """
    将JSON文件转换为Word文档

    Args:
        json_path: JSON文件路径
        docx_path: 输出Word文件路径（可选）

    Returns:
        生成的Word文件路径
    """
    json_path = Path(json_path)

    if not json_path.exists():
        raise FileNotFoundError(f"JSON文件不存在: {json_path}")

    # 读取JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # 默认输出路径
    if docx_path is None:
        docx_path = json_path.with_suffix('.docx')
    else:
        docx_path = Path(docx_path)

    # 创建文档
    doc = Document()

    # 设置默认字体（中文支持）
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    # 文档标题
    if 'title' in data:
        title = doc.add_heading(data['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 处理各章节
    for section in data.get('sections', []):
        # 章节标题
        if 'heading' in section:
            doc.add_heading(section['heading'], level=section.get('level', 1))

        section_type = section.get('type', 'text')

        if section_type == 'table':
            # 表格
            headers = section.get('headers', [])
            rows = section.get('rows', [])
            header_color = section.get('header_color', 'D9E2F3')

            if headers and rows:
                create_table(doc, headers, rows, header_color)
                doc.add_paragraph()  # 表格后空行

        elif section_type == 'text':
            # 段落文本
            content = section.get('content', '')
            if content:
                p = doc.add_paragraph(content)
                # 检查是否需要加粗部分文字
                if 'bold_prefix' in section:
                    # 重新创建带格式的段落
                    p.clear()
                    run_bold = p.add_run(section['bold_prefix'])
                    run_bold.bold = True
                    p.add_run(section.get('content_after_bold', ''))

        elif section_type == 'list':
            # 列表
            items = section.get('items', [])
            for item in items:
                doc.add_paragraph(item, style='List Bullet')

        elif section_type == 'key_value':
            # 键值对表格（两列）
            items = section.get('items', [])
            if items:
                headers = ['项目', '信息']
                rows = [[item['key'], item['value']] for item in items]
                header_color = section.get('header_color', 'D9E2F3')
                create_table(doc, headers, rows, header_color)
                doc.add_paragraph()

    # 保存文档
    doc.save(docx_path)
    print(f"Word文档已生成: {docx_path}")

    return str(docx_path)


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    json_path = sys.argv[1]
    docx_path = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        result = json_to_word(json_path, docx_path)
        print(f"转换成功: {result}")
    except Exception as e:
        print(f"转换失败: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
