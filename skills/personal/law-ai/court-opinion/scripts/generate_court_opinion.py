#!/usr/bin/env python3
"""
庭审代理意见书生成器 - 民事/刑事/劳动仲裁

用法:
    python3 generate_court_opinion.py --input court_opinion.json --output 代理意见书.docx

JSON 格式见 references/court_opinion_template.json
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


def add_legal_disclaimer(doc):
    """在文书末尾添加法条可靠性声明（灰色小字，不影响主体视觉效果）"""
    doc.add_paragraph()
    disclaimer_p = doc.add_paragraph()
    disclaimer_p.paragraph_format.space_before = Pt(6)
    disclaimer_p.paragraph_format.space_after = Pt(0)
    border_run = disclaimer_p.add_run('─' * 40)
    border_run.font.name = '宋体'
    border_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    border_run.font.size = Pt(9)
    border_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    disclaimer_text = (
        '【法条可靠性声明】本文书中引用的法条由 AI 辅助生成，'
        '可能存在版本滞后或条文号偏差。标注"[库中已核实]"的条文来自本地法条白名单，'
        '标注"[待核实]"的条文需人工核实。建议使用前登录国家法律法规数据库'
        '（flk.npc.gov.cn）核实相关法条的现行有效版本。'
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(disclaimer_text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def set_font(run, name='仿宋_GB2312', size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_page_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_section_title(doc, text, level='一'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, '黑体', 14, bold=True)
    return p


def add_subsection(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, '黑体', 12, bold=True)
    return p


def add_para(doc, text, indent=True):
    if not text:
        return
    for line in str(text).split('\n'):
        if line.strip():
            p = doc.add_paragraph()
            if indent:
                p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line.strip())
            set_font(run, '仿宋_GB2312', 12)


def generate_court_opinion(data, output_path):
    """生成庭审代理意见书"""
    doc = Document()
    set_page_margins(doc)

    doc.styles['Normal'].font.name = '仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    doc_type = data.get('文书类型', '民事代理意见书')

    # ===== 标题（黑体二号居中）=====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run(doc_type)
    set_font(title_run, '黑体', 22, bold=True)

    # ===== 代理人基本信息 =====
    case_name = data.get('案件名称', '')
    case_no = data.get('案号', '')
    party_side = data.get('代理方', '')
    lawyer = data.get('律师', '')
    lawyer_id = data.get('律师证号', '')
    law_firm = data.get('律所', '')
    court_date = data.get('庭审日期', datetime.now().strftime('%Y年%m月%d日'))

    info_data = [
        ['案件名称', case_name],
        ['案号', case_no],
        ['代理方', party_side],
        ['代理律师', lawyer],
        ['律师证号', lawyer_id],
        ['律师事务所', law_firm],
        ['庭审日期', court_date],
    ]
    info_table = doc.add_table(rows=len(info_data), cols=2)
    info_table.style = 'Table Grid'
    for i, (key, val) in enumerate(info_data):
        k_cell = info_table.rows[i].cells[0]
        v_cell = info_table.rows[i].cells[1]
        k_cell.text = key
        v_cell.text = val
        set_cell_bg(k_cell, 'D9E2F3')
        if k_cell.paragraphs[0].runs:
            k_cell.paragraphs[0].runs[0].font.bold = True
        k_cell.width = Cm(3.5)
        v_cell.width = Cm(13)

    doc.add_paragraph()

    # 致辞（敬语开头）
    court_name = data.get('法院', '尊敬的合议庭')
    intro_p = doc.add_paragraph()
    run = intro_p.add_run(f'{court_name}：')
    set_font(run, '仿宋_GB2312', 12)

    # 受托说明
    intro_text = data.get('受托说明',
        f'本代理人受{party_side}委托，依法担任本案诉讼代理人。'
        f'现根据庭审查明的事实、证据及相关法律规定，'
        f'发表如下代理意见，供合议庭参考：')
    add_para(doc, intro_text)
    doc.add_paragraph()

    # ===== 第一部分：事实意见 =====
    add_section_title(doc, '第一部分  事实意见')

    fact_data = data.get('事实意见', {})

    # 己方事实陈述
    own_facts = fact_data.get('己方事实陈述', '')
    if own_facts:
        add_subsection(doc, '一、己方事实陈述')
        add_para(doc, own_facts)

    # 对方主张反驳
    rebuttal = fact_data.get('对方主张反驳', '')
    if rebuttal:
        add_subsection(doc, '二、对对方主张的反驳')
        add_para(doc, rebuttal)

    doc.add_paragraph()

    # ===== 第二部分：证据意见 =====
    add_section_title(doc, '第二部分  证据意见')
    evidence_opinion = data.get('证据意见', '')
    if evidence_opinion:
        add_para(doc, evidence_opinion)
    doc.add_paragraph()

    # ===== 第三部分：法律适用分析 =====
    add_section_title(doc, '第三部分  法律适用分析')

    legal_analysis = data.get('法律适用', [])
    if isinstance(legal_analysis, list):
        for item in legal_analysis:
            if isinstance(item, dict):
                # 法条引用
                law_article = item.get('法条', '')
                law_content = item.get('内容', '')
                analysis = item.get('分析', '')

                if law_article:
                    quote_p = doc.add_paragraph()
                    quote_p.paragraph_format.left_indent = Cm(0.74)
                    quote_p.paragraph_format.space_after = Pt(4)
                    run = quote_p.add_run(f'【{law_article}】')
                    set_font(run, '黑体', 12, bold=True, color=RGBColor(0x17, 0x37, 0x6A))

                if law_content:
                    content_p = doc.add_paragraph()
                    content_p.paragraph_format.left_indent = Cm(1.48)
                    run = content_p.add_run(law_content)
                    set_font(run, '仿宋_GB2312', 11, color=RGBColor(0x44, 0x44, 0x44))

                if analysis:
                    add_para(doc, analysis)

                doc.add_paragraph()
            else:
                add_para(doc, str(item))
    else:
        add_para(doc, str(legal_analysis))

    # ===== 请求事项 =====
    add_section_title(doc, '请求事项')
    requests = data.get('请求事项', [])
    if requests:
        add_para(doc, '综上所述，本代理人请求法庭：', indent=True)
        for item in requests:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(item)
            set_font(run, '仿宋_GB2312', 12)

    doc.add_paragraph()

    # 结语
    closing = data.get('结语', '以上代理意见，请合议庭参考采纳。')
    add_para(doc, closing)

    doc.add_paragraph()

    # ===== 落款 =====
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_run = sign_p.add_run(
        f'代理律师：{lawyer}\n'
        f'律师证号：{lawyer_id}\n'
        f'律师事务所：{law_firm}\n'
        f'日期：{court_date}'
    )
    set_font(sign_run, '仿宋_GB2312', 12)

    # ===== 法条可靠性声明（灰色小字）=====
    add_legal_disclaimer(doc)

    doc.save(output_path)
    print(f"{doc_type}已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description='庭审代理意见书生成器',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument('--input', required=True, help='代理意见数据 JSON 文件路径')
    parser.add_argument('--output', default='代理意见书.docx', help='输出 Word 文件路径')

    args = parser.parse_args()

    json_path = Path(args.input)
    if not json_path.exists():
        print(f"错误: JSON 文件不存在: {json_path}")
        sys.exit(1)

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    try:
        generate_court_opinion(data, args.output)
    except Exception as e:
        print(f"生成失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
