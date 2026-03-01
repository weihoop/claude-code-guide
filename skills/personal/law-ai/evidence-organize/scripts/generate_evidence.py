#!/usr/bin/env python3
"""
证据目录 + 质证意见书生成器

用法:
    python3 generate_evidence.py --input 证据信息.json --output 证据目录及质证意见.docx

JSON 格式见 references/evidence_template.json
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


def set_font(run, name='仿宋_GB2312', size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_page_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def build_table(doc, headers, rows, header_color='2F5496', col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        set_cell_bg(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, '黑体', 11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # 数据行
    for row_idx, row_data in enumerate(rows):
        fill = 'EBF1F8' if row_idx % 2 == 0 else 'FFFFFF'
        for col_idx, val in enumerate(row_data):
            if col_idx >= len(table.rows[row_idx + 1].cells):
                continue
            cell = table.rows[row_idx + 1].cells[col_idx]
            set_cell_bg(cell, fill)
            p = cell.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                           if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(str(val))
            set_font(run, '仿宋_GB2312', 11)

    # 设置列宽
    if col_widths:
        for col_idx, width in enumerate(col_widths):
            for row in table.rows:
                if col_idx < len(row.cells):
                    row.cells[col_idx].width = Cm(width)

    return table


def generate_evidence(data, output_path):
    """生成证据目录 + 质证意见书"""
    doc = Document()
    set_page_margins(doc)

    doc.styles['Normal'].font.name = '仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # ===== 标题 =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run('证据目录及质证意见')
    set_font(run, '黑体', 22, bold=True)

    # 案件信息
    case_name = data.get('案件名称', '')
    case_no = data.get('案号', '')
    submitter = data.get('出具方', '')
    date = data.get('日期', datetime.now().strftime('%Y年%m月%d日'))

    for text in [case_name, f'案号：{case_no}', f'出具方：{submitter}', f'日期：{date}']:
        if text.strip('：'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_font(run, '仿宋_GB2312', 11)

    doc.add_paragraph()
    doc.add_page_break()

    # ===== 表1：己方证据目录 =====
    own_evidence = data.get('己方证据', [])
    if own_evidence:
        h1_p = doc.add_paragraph()
        h1_p.paragraph_format.space_before = Pt(12)
        h1_p.paragraph_format.space_after = Pt(8)
        run = h1_p.add_run('一、己方证据目录')
        set_font(run, '黑体', 14, bold=True)

        headers = ['序号', '证据名称', '证明目的', '证据来源', '页码']
        rows = []
        for ev in own_evidence:
            rows.append([
                str(ev.get('序号', '')),
                ev.get('证据名称', ''),
                ev.get('证明目的', ''),
                ev.get('来源', ''),
                str(ev.get('页码', '')),
            ])
        build_table(doc, headers, rows,
                    col_widths=[1.5, 4, 5, 3, 1.5])
        doc.add_paragraph()

    # ===== 表2：对方证据质证意见 =====
    opposite_evidence = data.get('对方证据', [])
    if opposite_evidence:
        h2_p = doc.add_paragraph()
        h2_p.paragraph_format.space_before = Pt(12)
        h2_p.paragraph_format.space_after = Pt(8)
        run = h2_p.add_run('二、对方证据质证意见')
        set_font(run, '黑体', 14, bold=True)

        # 三性说明
        note_p = doc.add_paragraph()
        run = note_p.add_run(
            '注：真实性——证据是否真实存在，有无伪造变造；'
            '合法性——取证方式是否合法；'
            '关联性——是否与本案待证事实有关。'
        )
        set_font(run, '仿宋_GB2312', 10)
        note_p.paragraph_format.space_after = Pt(6)

        headers = ['序号', '证据名称', '真实性', '合法性', '关联性', '质证意见']
        rows = []
        for ev in opposite_evidence:
            rows.append([
                str(ev.get('序号', '')),
                ev.get('证据名称', ''),
                ev.get('真实性', '无异议'),
                ev.get('合法性', '无异议'),
                ev.get('关联性', '无异议'),
                ev.get('质证意见', ''),
            ])
        build_table(doc, headers, rows,
                    col_widths=[1.5, 3.5, 1.5, 1.5, 1.5, 5.5])
        doc.add_paragraph()

    # ===== 落款 =====
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_run = sign_p.add_run(
        f'出具方：{submitter}\n'
        f'律师事务所：____________________\n'
        f'日期：{date}'
    )
    set_font(sign_run, '仿宋_GB2312', 12)

    doc.save(output_path)
    print(f"证据目录及质证意见已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description='证据目录 + 质证意见书生成器',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument('--input', required=True, help='证据信息 JSON 文件路径')
    parser.add_argument('--output', default='证据目录及质证意见.docx', help='输出 Word 文件路径')

    args = parser.parse_args()

    json_path = Path(args.input)
    if not json_path.exists():
        print(f"错误: JSON 文件不存在: {json_path}")
        sys.exit(1)

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    try:
        generate_evidence(data, args.output)
    except Exception as e:
        print(f"生成失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
