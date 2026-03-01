#!/usr/bin/env python3
"""
企业合规审查报告生成器

用法:
    python3 generate_compliance.py --area labor --input compliance_result.json --output 劳动合规报告.docx
    python3 generate_compliance.py --area data --input compliance_result.json --output 数据合规报告.docx
    python3 generate_compliance.py --area contract --input compliance_result.json --output 合同合规报告.docx
    python3 generate_compliance.py --area general --input compliance_result.json --output 综合合规报告.docx

审查领域:
    labor    - 劳动用工合规
    data     - 数据安全合规（个人信息保护）
    contract - 合同管理合规
    general  - 综合合规
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


AREA_NAMES = {
    'labor': '劳动用工合规',
    'data': '数据安全合规',
    'contract': '合同管理合规',
    'general': '综合合规',
}

RISK_CONFIG = {
    '高': {'fill': 'FFCCCC', 'color': RGBColor(0xC0, 0x00, 0x00), 'desc': '存在违法违规，需立即整改'},
    '中': {'fill': 'FFE5CC', 'color': RGBColor(0xFF, 0x66, 0x00), 'desc': '存在制度漏洞，建议3个月内完善'},
    '低': {'fill': 'FFFFCC', 'color': RGBColor(0xCC, 0x99, 0x00), 'desc': '流程不规范，建议逐步优化'},
}


def add_legal_disclaimer(doc):
    """在文书末尾添加法条可靠性声明（灰色小字，不影响主体视觉效果）"""
    doc.add_paragraph()
    disclaimer_p = doc.add_paragraph()
    disclaimer_p.paragraph_format.space_before = Pt(6)
    disclaimer_p.paragraph_format.space_after = Pt(0)
    border_run = disclaimer_p.add_run('─' * 40)
    border_run.font.name = '宋体'
    border_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    border_run.font.size = Pt(9)
    border_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    disclaimer_text = (
        '【法条可靠性声明】本文书中引用的法条由 AI 辅助生成，'
        '可能存在版本滞后或条文号偏差。标注"[库中已核实]"的条文来自本地法条白名单，'
        '标注"[待核实]"的条文需人工核实。建议使用前登录国家法律法规数据库'
        '（flk.npc.gov.cn）核实相关法条的现行有效版本。'
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(disclaimer_text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def set_font(run, name='仿宋_GB2312', size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_page_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, '黑体', 14, bold=True)
    return p


def add_para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_font(run, '仿宋_GB2312', 12)
    return p


def generate_compliance(data, area, output_path):
    """生成企业合规审查报告"""
    doc = Document()
    set_page_margins(doc)

    doc.styles['Normal'].font.name = '仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    area_name = AREA_NAMES.get(area, '合规')

    # ===== 标题 =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run(f'{area_name}审查报告')
    set_font(run, '黑体', 22, bold=True)

    # 企业名称副标题
    company = data.get('企业名称', '')
    if company:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_p.add_run(f'—— {company}')
        set_font(run, '黑体', 14)

    doc.add_paragraph()

    # ===== 一、审查概况 =====
    add_section_title(doc, '一、审查概况')

    info_data = [
        ['被审查单位', data.get('企业名称', '')],
        ['审查领域', area_name],
        ['审查范围', data.get('审查范围', '')],
        ['审查律师', data.get('审查律师', '')],
        ['律师事务所', data.get('律所', '')],
        ['审查日期', data.get('审查日期', datetime.now().strftime('%Y年%m月%d日'))],
    ]
    info_table = doc.add_table(rows=len(info_data), cols=2)
    info_table.style = 'Table Grid'
    for i, (key, val) in enumerate(info_data):
        k_cell = info_table.rows[i].cells[0]
        v_cell = info_table.rows[i].cells[1]
        k_cell.text = key
        v_cell.text = val
        set_cell_bg(k_cell, 'D9E2F3')
        if k_cell.paragraphs[0].runs:
            k_cell.paragraphs[0].runs[0].font.bold = True
        k_cell.width = Cm(3.5)
        v_cell.width = Cm(13)
    doc.add_paragraph()

    # 整体评价
    overall = data.get('整体评价', '')
    if overall:
        add_para(doc, overall)
    doc.add_paragraph()

    # ===== 二、合规问题清单 =====
    add_section_title(doc, '二、合规问题清单')

    # 风险等级说明
    for level, config in RISK_CONFIG.items():
        note_p = doc.add_paragraph()
        note_p.paragraph_format.space_after = Pt(3)
        run = note_p.add_run(f'  {level}风险：{config["desc"]}')
        run.font.color.rgb = config['color']
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run.font.size = Pt(11)
    doc.add_paragraph()

    issues = data.get('问题清单', [])
    if issues:
        headers = ['序号', '问题描述', '风险等级', '法律依据', '整改建议']
        table = doc.add_table(rows=1 + len(issues), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            set_cell_bg(cell, '2F5496')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            set_font(run, '黑体', 11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

        # 数据行
        for idx, issue in enumerate(issues):
            level = issue.get('风险等级', '低')
            config = RISK_CONFIG.get(level, RISK_CONFIG['低'])
            row = table.rows[idx + 1].cells

            values = [
                str(issue.get('序号', idx + 1)),
                issue.get('问题描述', ''),
                level,
                issue.get('法律依据', ''),
                issue.get('整改建议', ''),
            ]
            for col, val in enumerate(values):
                row[col].text = val
                row[col].paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if col in [0, 2] else WD_ALIGN_PARAGRAPH.LEFT
                )
                if col == 2:
                    set_cell_bg(row[col], config['fill'])
                    if row[col].paragraphs[0].runs:
                        row[col].paragraphs[0].runs[0].font.color.rgb = config['color']
                        row[col].paragraphs[0].runs[0].font.bold = True

        # 列宽
        col_widths = [1.2, 5.5, 1.5, 3.5, 5.5]
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)

    doc.add_paragraph()

    # ===== 三、整改建议（按优先级） =====
    add_section_title(doc, '三、整改建议')

    # 按风险等级分组输出
    high_issues = [i for i in issues if i.get('风险等级') == '高']
    mid_issues = [i for i in issues if i.get('风险等级') == '中']
    low_issues = [i for i in issues if i.get('风险等级') == '低']

    for level_name, level_issues in [('高风险（立即整改）', high_issues),
                                       ('中风险（近期整改）', mid_issues),
                                       ('低风险（逐步优化）', low_issues)]:
        if level_issues:
            p = doc.add_paragraph()
            run = p.add_run(f'【{level_name}】')
            set_font(run, '黑体', 12, bold=True)
            p.paragraph_format.space_before = Pt(8)
            for item in level_issues:
                suggestion = item.get('整改建议', '')
                if suggestion:
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.left_indent = Cm(0.74)
                    p2.paragraph_format.space_after = Pt(4)
                    run = p2.add_run(f'• {suggestion}')
                    set_font(run, '仿宋_GB2312', 12)

    doc.add_paragraph()

    # ===== 四、法律依据 =====
    add_section_title(doc, '四、法律依据')
    legal_bases = set()
    for issue in issues:
        lb = issue.get('法律依据', '')
        if lb:
            legal_bases.add(lb)

    if legal_bases:
        for lb in sorted(legal_bases):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f'• {lb}')
            set_font(run, '仿宋_GB2312', 12)
    else:
        add_para(doc, '（法律依据详见各问题整改建议）')

    doc.add_paragraph()

    # ===== 落款 =====
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lawyer = data.get('审查律师', '____________________')
    law_firm = data.get('律所', '____________________')
    date = data.get('审查日期', datetime.now().strftime('%Y年%m月%d日'))
    sign_run = sign_p.add_run(
        f'审查律师：{lawyer}\n'
        f'律师事务所：{law_firm}\n'
        f'出具日期：{date}'
    )
    set_font(sign_run, '仿宋_GB2312', 12)

    # ===== 法条可靠性声明（灰色小字）=====
    add_legal_disclaimer(doc)

    doc.save(output_path)
    print(f"合规审查报告已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description='企业合规审查报告生成器',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument('--area', required=True,
                        choices=['labor', 'data', 'contract', 'general'],
                        help='审查领域')
    parser.add_argument('--input', required=True, help='审查结果 JSON 文件路径')
    parser.add_argument('--output', default='合规审查报告.docx', help='输出 Word 文件路径')

    args = parser.parse_args()

    json_path = Path(args.input)
    if not json_path.exists():
        print(f"错误: JSON 文件不存在: {json_path}")
        sys.exit(1)

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    try:
        generate_compliance(data, args.area, args.output)
    except Exception as e:
        print(f"生成失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
