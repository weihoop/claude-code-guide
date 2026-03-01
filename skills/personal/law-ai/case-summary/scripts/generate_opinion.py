#!/usr/bin/env python3
"""
法律意见书生成器

用法:
    python3 generate_opinion.py --input opinion_data.json --output 法律意见书.docx

JSON 格式见 references/opinion_template.json
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


def add_legal_disclaimer(doc):
    """在文书末尾添加法条可靠性声明（灰色小字，不影响主体视觉效果）"""
    doc.add_paragraph()
    disclaimer_p = doc.add_paragraph()
    disclaimer_p.paragraph_format.space_before = Pt(6)
    disclaimer_p.paragraph_format.space_after = Pt(0)
    # 分隔线
    border_run = disclaimer_p.add_run('─' * 40)
    border_run.font.name = '宋体'
    border_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    border_run.font.size = Pt(9)
    border_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    disclaimer_text = (
        '【法条可靠性声明】本文书中引用的法条由 AI 辅助生成，'
        '可能存在版本滞后或条文号偏差。标注"[库中已核实]"的条文来自本地法条白名单，'
        '标注"[待核实]"的条文需人工核实。建议使用前登录国家法律法规数据库'
        '（flk.npc.gov.cn）核实相关法条的现行有效版本。'
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(disclaimer_text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def set_font(run, name='仿宋_GB2312', size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_page_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_section(doc, num, title, content_func, *args):
    """添加编号章节"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'{num}、{title}')
    set_font(run, '黑体', 14, bold=True)
    content_func(doc, *args)


def add_text_block(doc, text):
    """添加文本段落"""
    if not text:
        return
    for para in str(text).split('\n'):
        if para.strip():
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(para.strip())
            set_font(run, '仿宋_GB2312', 12)


def add_numbered_list(doc, items):
    """添加编号列表"""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f'{i}. {item}')
        set_font(run, '仿宋_GB2312', 12)


def add_bullet_list(doc, items):
    """添加项目符号列表"""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f'• {item}')
        set_font(run, '仿宋_GB2312', 12)


def generate_opinion(data, output_path):
    """生成法律意见书"""
    doc = Document()
    set_page_margins(doc)

    doc.styles['Normal'].font.name = '仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # ===== 标题 =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(12)
    run = title_p.add_run('法律意见书')
    set_font(run, '黑体', 22, bold=True)

    # ===== 一、委托方基本信息 =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('一、委托方基本信息')
    set_font(run, '黑体', 14, bold=True)

    # 基本信息表格
    info_fields = [
        ('委托方', data.get('委托方', '')),
        ('案件类型', data.get('案件类型', '')),
        ('受托律师', data.get('律师', '')),
        ('律师事务所', data.get('律所', '')),
        ('受托日期', data.get('受托日期', datetime.now().strftime('%Y年%m月%d日'))),
    ]
    info_table = doc.add_table(rows=len(info_fields), cols=2)
    info_table.style = 'Table Grid'
    for i, (key, val) in enumerate(info_fields):
        key_cell = info_table.rows[i].cells[0]
        val_cell = info_table.rows[i].cells[1]
        key_cell.text = key
        val_cell.text = val
        set_cell_bg(key_cell, 'D9E2F3')
        key_cell.paragraphs[0].runs[0].font.bold = True
        key_cell.width = Cm(3.5)
        val_cell.width = Cm(13)
    doc.add_paragraph()

    # ===== 二、案件事实概述 =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('二、案件事实概述')
    set_font(run, '黑体', 14, bold=True)
    add_text_block(doc, data.get('案件事实', ''))
    doc.add_paragraph()

    # ===== 三、法律关系分析 =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('三、法律关系分析')
    set_font(run, '黑体', 14, bold=True)
    add_text_block(doc, data.get('法律关系分析', ''))
    doc.add_paragraph()

    # ===== 四、争议焦点 =====
    dispute_points = data.get('争议焦点', [])
    if dispute_points:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run('四、争议焦点')
        set_font(run, '黑体', 14, bold=True)
        if isinstance(dispute_points, list):
            add_bullet_list(doc, dispute_points)
        else:
            add_text_block(doc, dispute_points)
        doc.add_paragraph()

    # ===== 五、相关法律规定 =====
    legal_rules = data.get('法律规定', [])
    if legal_rules:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run('五、相关法律规定')
        set_font(run, '黑体', 14, bold=True)
        if isinstance(legal_rules, list):
            for rule in legal_rules:
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Cm(0.74)
                p2.paragraph_format.space_after = Pt(6)
                run = p2.add_run(rule)
                set_font(run, '仿宋_GB2312', 11)
        else:
            add_text_block(doc, legal_rules)
        doc.add_paragraph()

    # ===== 六、法律意见 =====
    opinions = data.get('法律意见', [])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('六、法律意见')
    set_font(run, '黑体', 14, bold=True)
    if isinstance(opinions, list):
        for opinion in opinions:
            add_text_block(doc, opinion)
    else:
        add_text_block(doc, opinions)
    doc.add_paragraph()

    # ===== 七、证据建议清单 =====
    evidence_list = data.get('证据建议', [])
    if evidence_list:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run('七、证据建议清单')
        set_font(run, '黑体', 14, bold=True)
        add_numbered_list(doc, evidence_list)
        doc.add_paragraph()

    # ===== 八、风险提示 =====
    risk_hint = data.get('风险提示', '')
    if risk_hint:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run('八、风险提示')
        set_font(run, '黑体', 14, bold=True)
        add_text_block(doc, risk_hint)
        doc.add_paragraph()

    # ===== 九、律师声明 =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('九、律师声明')
    set_font(run, '黑体', 14, bold=True)

    declaration = data.get('律师声明',
        '本法律意见书依据委托方提供的资料及现行有效的中国法律法规出具，'
        '仅供委托方参考使用，不得用于其他目的。本所对委托方提供资料的真实性、'
        '完整性不作独立核实，如委托方提供的资料存在虚假，本所不承担相应责任。')
    add_text_block(doc, declaration)
    doc.add_paragraph()

    # ===== 落款 =====
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lawyer = data.get('律师', '____________________')
    law_firm = data.get('律所', '____________________')
    date = data.get('受托日期', datetime.now().strftime('%Y年%m月%d日'))
    sign_run = sign_p.add_run(
        f'经办律师：{lawyer}\n'
        f'律师事务所：{law_firm}\n'
        f'日期：{date}'
    )
    set_font(sign_run, '仿宋_GB2312', 12)

    # ===== 法条可靠性声明（灰色小字）=====
    add_legal_disclaimer(doc)

    doc.save(output_path)
    print(f"法律意见书已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description='法律意见书生成器',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument('--input', required=True, help='意见书数据 JSON 文件路径')
    parser.add_argument('--output', default='法律意见书.docx', help='输出 Word 文件路径')

    args = parser.parse_args()

    json_path = Path(args.input)
    if not json_path.exists():
        print(f"错误: JSON 文件不存在: {json_path}")
        sys.exit(1)

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    try:
        generate_opinion(data, args.output)
    except Exception as e:
        print(f"生成失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
