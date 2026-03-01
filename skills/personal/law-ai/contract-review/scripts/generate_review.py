#!/usr/bin/env python3
"""
合同审查意见书生成器

用法:
    python3 generate_review.py --input review_result.json --output 合同审查意见书.docx

JSON 格式:
{
  "合同名称": "××买卖合同",
  "审查立场": "甲方",
  "合同类型": "买卖合同",
  "审查日期": "2025年XX月XX日",
  "总体评价": "该合同整体结构完整，但存在以下主要风险...",
  "风险清单": [
    {
      "条款位置": "第3条第2款",
      "风险等级": "高",
      "问题描述": "付款条件约定不明确",
      "修改建议": "建议明确约定付款时间节点和方式"
    }
  ],
  "缺失条款": ["违约责任条款", "争议解决条款"],
  "综合建议": "建议重点修改付款条款后再签署"
}
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


# 风险等级颜色映射
RISK_COLORS = {
    '高': RGBColor(0xC0, 0x00, 0x00),   # 深红
    '中': RGBColor(0xFF, 0x66, 0x00),   # 橙色
    '低': RGBColor(0xFF, 0xC0, 0x00),   # 黄色
}
RISK_FILLS = {
    '高': 'FFCCCC',
    '中': 'FFE5CC',
    '低': 'FFFFCC',
}


def add_legal_disclaimer(doc):
    """在文书末尾添加法条可靠性声明（灰色小字，不影响主体视觉效果）"""
    doc.add_paragraph()
    disclaimer_p = doc.add_paragraph()
    disclaimer_p.paragraph_format.space_before = Pt(6)
    disclaimer_p.paragraph_format.space_after = Pt(0)
    border_run = disclaimer_p.add_run('─' * 40)
    border_run.font.name = '宋体'
    border_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    border_run.font.size = Pt(9)
    border_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    disclaimer_text = (
        '【法条可靠性声明】本文书中引用的法条由 AI 辅助生成，'
        '可能存在版本滞后或条文号偏差。标注"[库中已核实]"的条文来自本地法条白名单，'
        '标注"[待核实]"的条文需人工核实。建议使用前登录国家法律法规数据库'
        '（flk.npc.gov.cn）核实相关法条的现行有效版本。'
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(disclaimer_text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def set_font(run, name='仿宋_GB2312', size=12, bold=False, color=None):
    """统一设置字体"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_bg(cell, color_hex):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_page_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    """设置页边距（cm）"""
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def add_para(doc, text, style='Normal', align=WD_ALIGN_PARAGRAPH.LEFT,
             font_name='仿宋_GB2312', size=12, bold=False, space_before=6, space_after=6):
    """添加段落"""
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, font_name, size, bold)
    return p


def generate_review(data, output_path):
    """生成合同审查意见书 Word 文档"""
    doc = Document()
    set_page_margins(doc)

    # 设置默认样式
    doc.styles['Normal'].font.name = '仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # ===== 封面 =====
    # 标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('合同审查意见书')
    set_font(title_run, '黑体', 22, bold=True)
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(12)

    # 副标题（合同名称）
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(f"——{data.get('合同名称', '')} 审查")
    set_font(subtitle_run, '黑体', 14)

    doc.add_paragraph()

    # 基本信息表
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ['合同名称', data.get('合同名称', '')],
        ['审查立场', data.get('审查立场', '')],
        ['合同类型', data.get('合同类型', '')],
        ['审查日期', data.get('审查日期', datetime.now().strftime('%Y年%m月%d日'))],
    ]
    for i, (key, val) in enumerate(info_data):
        key_cell = info_table.rows[i].cells[0]
        val_cell = info_table.rows[i].cells[1]
        key_cell.text = key
        val_cell.text = val
        key_cell.paragraphs[0].runs[0].font.bold = True
        set_cell_bg(key_cell, 'D9E2F3')

    doc.add_paragraph()
    doc.add_page_break()

    # ===== 一、总体评价 =====
    add_para(doc, '一、总体评价', font_name='黑体', size=14, bold=True, space_before=12)
    add_para(doc, data.get('总体评价', ''), size=12)

    # ===== 二、风险清单 =====
    add_para(doc, '二、风险清单', font_name='黑体', size=14, bold=True, space_before=12)

    risk_list = data.get('风险清单', [])
    if risk_list:
        headers = ['序号', '条款位置', '风险等级', '问题描述', '修改建议']
        risk_table = doc.add_table(rows=1 + len(risk_list), cols=len(headers))
        risk_table.style = 'Table Grid'
        risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        header_row = risk_table.rows[0].cells
        for i, h in enumerate(headers):
            header_row[i].text = h
            set_cell_bg(header_row[i], '2F5496')
            p = header_row[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run(h)
            run.text = h
            set_font(run, '黑体', 11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

        # 数据行
        for idx, item in enumerate(risk_list):
            row = risk_table.rows[idx + 1].cells
            level = item.get('风险等级', '低')
            fill_color = RISK_FILLS.get(level, 'FFFFFF')

            values = [
                str(idx + 1),
                item.get('条款位置', ''),
                level,
                item.get('问题描述', ''),
                item.get('修改建议', ''),
            ]
            for col, val in enumerate(values):
                row[col].text = val
                row[col].paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if col in [0, 2] else WD_ALIGN_PARAGRAPH.LEFT
                )
                if col == 2:  # 风险等级列着色
                    set_cell_bg(row[col], fill_color)
                    run = row[col].paragraphs[0].runs
                    if run:
                        run[0].font.color.rgb = RISK_COLORS.get(level, RGBColor(0, 0, 0))
                        run[0].font.bold = True

        # 设置列宽
        col_widths = [1, 3, 2, 6, 6]
        for col_idx, width in enumerate(col_widths):
            for row in risk_table.rows:
                row.cells[col_idx].width = Cm(width)

    doc.add_paragraph()

    # ===== 三、缺失条款 =====
    missing = data.get('缺失条款', [])
    if missing:
        add_para(doc, '三、缺失条款说明', font_name='黑体', size=14, bold=True, space_before=12)
        add_para(doc, '以下重要条款在合同中未作约定，建议补充：', size=12)
        for item in missing:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(item)
            set_font(run, '仿宋_GB2312', 12)

    doc.add_paragraph()

    # ===== 四、综合建议 =====
    add_para(doc, '四、综合建议', font_name='黑体', size=14, bold=True, space_before=12)
    add_para(doc, data.get('综合建议', ''), size=12)

    doc.add_paragraph()

    # ===== 落款 =====
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_run = sign_p.add_run(
        f"审查律师：____________________\n"
        f"律师事务所：____________________\n"
        f"日期：{data.get('审查日期', datetime.now().strftime('%Y年%m月%d日'))}"
    )
    set_font(sign_run, '仿宋_GB2312', 12)

    # ===== 法条可靠性声明（灰色小字）=====
    add_legal_disclaimer(doc)

    doc.save(output_path)
    print(f"合同审查意见书已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description='合同审查意见书生成器',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument('--input', required=True, help='审查结果 JSON 文件路径')
    parser.add_argument('--output', default='合同审查意见书.docx', help='输出 Word 文件路径')

    args = parser.parse_args()

    json_path = Path(args.input)
    if not json_path.exists():
        print(f"错误: JSON 文件不存在: {json_path}")
        sys.exit(1)

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    try:
        generate_review(data, args.output)
    except Exception as e:
        print(f"生成失败: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
